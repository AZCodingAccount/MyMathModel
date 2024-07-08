import pandas as pd
from sklearn.metrics import r2_score
from sklearn.model_selection import train_test_split
from sklearn.linear_model import LinearRegression
from sklearn.pipeline import make_pipeline
from sklearn.preprocessing import StandardScaler, PolynomialFeatures
from docx import Document
import numpy as np

# TODO:进行线性,多项式，指数，幂函数的回归，x值为平均销售单价，y值为销售量
# 加载数据
file_path = "../第二问结果/第二问最终数据.xlsx"
df = pd.read_excel(file_path, header=[0, 1])

# 数据预处理
df = df.drop([0, 2], axis=0)
df = df.drop(df.columns[0], axis=1)
df.fillna(0, inplace=True)
df.columns = ['_'.join(map(str, col)) for col in df.columns]

# 创建模型的pipeline
poly_reg = make_pipeline(PolynomialFeatures(degree=2), LinearRegression())
exp_reg = make_pipeline(StandardScaler(), LinearRegression())  # 指数回归
power_reg = make_pipeline(StandardScaler(), PolynomialFeatures(degree=1), LinearRegression())  # 幂函数回归

models = {
    "Linear Regression": LinearRegression(),
    "Polynomial Regression": poly_reg,
    "Exponential Regression": exp_reg,
    "Power Regression": power_reg,
}

# 初始化文档
doc = Document()
doc.add_heading('模型分析报告', 0)

for i in range(0, len(df.columns) - 1, 6):
    product_type = df.columns[i].split('_')[1]

    X = df[[df.columns[i + 1]]]  # 只选择平均销售单价作为特征
    y = df[df.columns[i]]

    # 剔除X为0的行，并相应地剔除y中对应的行
    mask = (X.iloc[:, 0] != 0)
    X = X[mask]
    y = y[mask]

    # 之后检查y，如果y为0，再剔除X中的对应行
    mask_y = (y != 0)
    X = X[mask_y]
    y = y[mask_y]

    # 如果在剔除0之后X或y为空，则继续下一个循环
    if X.empty or y.empty:
        continue

    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

    for model_name, model in models.items():
        if model_name == "Exponential Regression":
            y_train_log = np.log(y_train + 1e-5)  # 避免log(0)
            y_test_log = np.log(y_test + 1e-5)
            model.fit(X_train, y_train_log)
            y_pred_log = model.predict(X_test)
            y_pred = np.exp(y_pred_log) - 1e-5
            r2 = r2_score(y_test, y_pred)
        elif model_name == "Power Regression":
            X_train_log = np.log(X_train + 1e-5)
            X_test_log = np.log(X_test + 1e-5)
            y_train_log = np.log(y_train + 1e-5)
            model.fit(X_train_log, y_train_log)
            y_pred_log = model.predict(X_test_log)
            y_pred = np.exp(y_pred_log) - 1e-5
            r2 = r2_score(y_test, y_pred)
        else:
            model.fit(X_train, y_train)
            y_pred = model.predict(X_test)
            r2 = r2_score(y_test, y_pred)

        # 下面的代码将结果添加到文档中
        doc.add_heading(f"对于 {product_type} 的 {model_name}:", level=1)

        # 根据模型类型记录相关信息
        if model_name == "Linear Regression":
            beta_0 = model.intercept_
            beta_1 = model.coef_[0]
            doc.add_paragraph(f"模型为：y = {beta_0:.4f} + {beta_1:.4f} * 平均销售单价")

        elif model_name == "Polynomial Regression":
            poly_features = model.named_steps['polynomialfeatures']
            lin_reg = model.named_steps['linearregression']
            coefficients = lin_reg.coef_
            intercept = lin_reg.intercept_
            doc.add_paragraph(f"模型为：y = {intercept:.4f} + {coefficients[1]:.4f} * 平均销售单价 + "
                              f"{coefficients[2]:.4f} * 平均销售单价^2")

        elif model_name == "Exponential Regression":
            lin_reg = model.named_steps['linearregression']
            doc.add_paragraph(f"模型为：y = e^({lin_reg.intercept_:.4f} + {lin_reg.coef_[0]:.4f} * 平均销售单价)")

        elif model_name == "Power Regression":
            lin_reg = model.named_steps['linearregression']
            doc.add_paragraph(f"模型为：y = {np.exp(lin_reg.intercept_):.4f} * 平均销售单价^{lin_reg.coef_[1]:.4f}")

        doc.add_paragraph(f"R^2 分数为: {r2:.4f}")
        doc.add_paragraph("-" * 50)

doc.save("../第二问结果/模型分析报告.docx")
